"""
Générateur du Cahier de Test ISIMM — python-docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Background color for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_col_width(table, col_index, width_cm):
    for row in table.rows:
        row.cells[col_index].width = Cm(width_cm)

def add_heading(doc, text, level=1, color_hex="1F3864"):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16)
        )
    return p

def add_paragraph(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1E, 0x8B, 0x4C)
    p.paragraph_format.left_indent = Inches(0.3)
    # add light gray background via XML shading on paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p

HEADER_BG   = "1F3864"   # dark navy — header rows
HEADER_FG   = "FFFFFF"
SUB_BG      = "2E74B5"   # medium blue — sub-headers
ALT_ROW     = "EBF3FB"   # light blue — alternating rows
WHITE       = "FFFFFF"
GREEN_BG    = "E2EFDA"
ORANGE_BG   = "FCE4D6"
YELLOW_BG   = "FFF2CC"

def make_table(doc, headers, rows, col_widths=None,
               header_bg=HEADER_BG, alt_bg=ALT_ROW):
    """Create a styled table. headers = list of strings, rows = list of lists."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)

    # data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        bg = alt_bg if r_idx % 2 == 1 else WHITE
        for c_idx, cell_text in enumerate(row_data):
            cell = row_cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(8.5)

    # column widths
    if col_widths:
        for c_idx, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[c_idx].width = Cm(w)

    doc.add_paragraph()  # spacing
    return table


# ── document setup ────────────────────────────────────────────────────────────

doc = Document()

# Page margins
sections = doc.sections
for section in sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(2)

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ── COVER PAGE ────────────────────────────────────────────────────────────────

doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("CAHIER DE TEST & PLAN DE VALIDATION")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run("Plateforme ISIMM — Gestion des Préinscriptions 2025-2026")
run2.bold = True
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()

meta_data = [
    ("Projet",     "Plateforme de Préinscription ISIMM"),
    ("Version",    "1.0"),
    ("Date",       datetime.date.today().strftime("%d/%m/%Y")),
    ("Auteur",     "Donia Jellali"),
    ("Stack",      "Angular 17 · Django · PostgreSQL · PaddleOCR · ReportLab"),
    ("Parcours",   "MPGL · MPDS · MP3I · MRGL · MRMI · ING-GL"),
]
meta_tbl = doc.add_table(rows=len(meta_data), cols=2)
meta_tbl.style = 'Table Grid'
meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(meta_data):
    cells = meta_tbl.rows[i].cells
    set_cell_bg(cells[0], "1F3864")
    r = cells[0].paragraphs[0].add_run(k)
    r.bold = True; r.font.color.rgb = RGBColor(255,255,255); r.font.size = Pt(10)
    set_cell_bg(cells[1], "EBF3FB" if i % 2 == 0 else "FFFFFF")
    r2 = cells[1].paragraphs[0].add_run(v)
    r2.font.size = Pt(10)
    cells[0].width = Cm(5); cells[1].width = Cm(10)

doc.add_page_break()

# ── TABLE OF CONTENTS placeholder ─────────────────────────────────────────────

add_heading(doc, "TABLE DES MATIÈRES", 1)
toc_items = [
    "SECTION 0 — Environnement & Prérequis ......................................... 3",
    "SECTION 1 — Modules Techniques Critiques ....................................... 4",
    "  1.1  OCR (PaddleOCR) ......................................................... 4",
    "  1.2  Générateur PDF (ReportLab) ............................................... 6",
    "  1.3  Sécurité QR Code (SHA-256) .............................................. 8",
    "SECTION 2 — Parcours End-to-End par User Story ................................ 10",
    "  US-01  Admin configure les parcours .......................................... 10",
    "  US-02  Responsable synchronise et ouvre les offres ........................... 11",
    "  US-03  Candidat 1 postule à MRGL ............................................. 12",
    "  US-04  Candidat 1 dépose son dossier ......................................... 14",
    "  US-05  Candidat 2 postule à MP3I ............................................. 15",
    "  US-06  Responsable traite les candidatures ................................... 15",
    "  US-07  Candidat 1 suit sa candidature et télécharge le PDF .................. 16",
    "SECTION 3 — Tests de Régression Critiques ...................................... 17",
    "SECTION 4 — Matrice de Validation Finale ....................................... 18",
    "SECTION 5 — Commandes de Vérification DB Rapides .............................. 19",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Inches(0.2) if item.startswith("  ") else Inches(0)
    p.runs[0].font.size = Pt(10)

doc.add_page_break()

# ── SECTION 0 ─────────────────────────────────────────────────────────────────

add_heading(doc, "SECTION 0 — Environnement & Prérequis", 1)
add_paragraph(doc, "Avant de lancer les tests, vérifier que tous les services sont actifs et que les données de test sont en place.", italic=True)
doc.add_paragraph()

make_table(doc,
    ["Élément", "Valeur attendue"],
    [
        ["Frontend Angular", "http://localhost:4200 — ng serve actif, aucune erreur de compilation"],
        ["Backend Django", "http://localhost:8000 — tous les microservices actifs"],
        ["Base de données", "PostgreSQL/SQLite — migrations appliquées (python manage.py migrate)"],
        ["OCR Service", "PaddleOCR initialisé, paddle importable sans erreur"],
        ["Utilisateurs de test", "Admin, Responsable, 2 Candidats créés en DB"],
        ["Documents de test", "5 fichiers PDF/JPG préparés selon les scénarios OCR (DOC-01 à DOC-05)"],
    ],
    col_widths=[5, 13]
)

doc.add_page_break()

# ── SECTION 1.1 — OCR ─────────────────────────────────────────────────────────

add_heading(doc, "SECTION 1 — Modules Techniques Critiques", 1)
add_heading(doc, "1.1 — OCR (PaddleOCR)", 2)

add_paragraph(doc, "Endpoint cible : POST /api/ocr/analyser/ ou POST /api/ocr/analyser-lot/", bold=True)
doc.add_paragraph()

add_heading(doc, "Jeux de documents de test à préparer", 3, "2E74B5")
make_table(doc,
    ["ID Doc", "Type", "Format", "État", "Objectif"],
    [
        ["DOC-01", "Relevé de notes (scan propre)", "PDF, 300 DPI", "Lisible", "Cas nominal — taux réussite élevé"],
        ["DOC-02", "Diplôme licence (photo téléphone)", "JPG, 72 DPI", "Légèrement flou", "Cas dégradé — taux partiel"],
        ["DOC-03", "Fichier texte renommé en PDF", "PDF", "Non-image", "Cas erreur — document illisible"],
        ["DOC-04", "PDF protégé par mot de passe", "PDF", "Verrouillé", "Cas erreur — accès refusé"],
        ["DOC-05", "Relevé de notes (scan recto-verso)", "PDF multi-pages", "Lisible", "Cas multipage"],
    ],
    col_widths=[2, 4, 3, 3, 6]
)

add_heading(doc, "Scénarios de test OCR", 3, "2E74B5")
make_table(doc,
    ["ID", "Scénario", "Inputs", "Actions", "Résultat attendu Angular", "Résultat attendu DB"],
    [
        ["OCR-01", "Extraction nominale — relevé propre",
         "DOC-01 (PDF 300 DPI, relevé ISIMM)",
         "1. Candidat ouvre \"Dossiers de candidature\"\n2. Glisse-dépose DOC-01 sur la zone upload\n3. Cliquer \"Envoyer\"",
         "Barre de progression 0→100% ; message ✅ Document analysé ; champs M1/M2/M3 pré-remplis",
         "statut_ocr='succes' ; taux_confiance >= 0.75 ; ocr_data JSON rempli"],
        ["OCR-02", "Extraction dégradée — photo floue",
         "DOC-02 (JPG 72 DPI)",
         "Même flux que OCR-01 avec DOC-02",
         "Message ⚠️ Extraction partielle ; certains champs vides ; formulaire modifiable",
         "statut_ocr='partiel' ; taux_confiance entre 0.40 et 0.74 ; ocr_warnings non vide"],
        ["OCR-03", "Document illisible",
         "DOC-03 (fichier texte en PDF)",
         "Même flux avec DOC-03",
         "Message ❌ Impossible d'extraire les données ; bouton Réessayer visible",
         "statut_ocr='echec' ; taux_confiance=0 ; log d'erreur enregistré"],
        ["OCR-04", "PDF protégé",
         "DOC-04",
         "Même flux avec DOC-04",
         "Message ❌ Document protégé ou illisible immédiatement",
         "HTTP 422 ; statut_ocr='erreur_acces'"],
        ["OCR-05", "PDF multi-pages",
         "DOC-05 (2 pages)",
         "Même flux avec DOC-05",
         "Toutes les pages traitées ; données des 2 pages fusionnées ; champs cohérents",
         "pages_traitees=2 dans le JSON OCR ; données complètes"],
        ["OCR-06", "Charge — 3 docs simultanés",
         "DOC-01 + DOC-02 + DOC-05",
         "Upload 3 documents en séquence rapide (<5 sec d'intervalle)",
         "Chaque upload traité indépendamment ; aucun mélange de données",
         "3 entrées distinctes en DB avec leurs candidature_id respectifs"],
    ],
    col_widths=[1.5, 3, 3, 4, 4, 4]
)

add_heading(doc, "Script de vérification OCR (Django Shell)", 3, "2E74B5")
add_code_block(doc,
"""from dossier_service.models import DocumentDossier
d = DocumentDossier.objects.last()
print(d.statut_ocr, d.taux_confiance, d.ocr_data)
# Résultat attendu OCR-01 : ('succes', 0.87, {'M1': 13.5, 'M2': 14.2, ...})""")

doc.add_page_break()

# ── SECTION 1.2 — PDF ─────────────────────────────────────────────────────────

add_heading(doc, "1.2 — Générateur PDF (ReportLab)", 2)
add_paragraph(doc, "Endpoint cible : GET /api/candidatures/{id}/recu-preinscription/", bold=True)
doc.add_paragraph()

make_table(doc,
    ["ID", "Scénario", "Inputs", "Actions", "Résultat attendu PDF", "Résultat attendu Django"],
    [
        ["PDF-01", "Génération nominale — Master",
         "Candidature ID valide, statut soumis, master MRGL",
         "1. Connecté en Candidat\n2. Aller dans \"Candidatures\"\n3. Cliquer \"Récapitulatif\"\n4. Cliquer \"Télécharger PDF\"",
         "PDF téléchargé ; contient : Logo ISIMM, Nom candidat, N° candidature, Parcours, Date soumission, QR Code",
         "HTTP 200 ; Content-Type: application/pdf ; fichier >= 50 Ko"],
        ["PDF-02", "Vérification mise en page",
         "PDF issu de PDF-01",
         "Ouvrir le PDF dans un lecteur PDF",
         "Page A4 ; marges respectées ; police lisible (>=10pt) ; logo ISIMM visible ; QR Code >=2cm x 2cm",
         "—"],
        ["PDF-03", "Intégration données candidat",
         "Candidat \"Ahmed Ben Ali\", CIN \"12345678\", email \"ahmed@test.tn\"",
         "Générer le PDF après soumission de la candidature",
         "PDF contient exactement : \"Ahmed Ben Ali\", \"12345678\", \"ahmed@test.tn\", date JJ/MM/AAAA",
         "Données identiques à Candidature.objects.get(id=X)"],
        ["PDF-04", "Données manquantes — champs vides",
         "Candidat dont phone et address sont vides",
         "Générer le PDF",
         "PDF généré sans erreur ; champs vides affichés comme \"—\" ou \"Non renseigné\" ; pas de crash NoneType",
         "HTTP 200 ; aucune exception dans les logs Django"],
        ["PDF-05", "2 candidatures — PDFs distincts",
         "Candidat avec 2 candidatures (MRGL + MPGL)",
         "Générer PDF pour chaque candidature séparément",
         "Chaque PDF contient le bon N° de candidature et le bon parcours ; aucune confusion",
         "2 fichiers différents ; N° candidature différents"],
        ["PDF-06", "Charge — 3 clics rapides",
         "1 compte Candidat",
         "Cliquer 3 fois \"Télécharger\" en 2 secondes",
         "1 seul PDF téléchargé ou 3 identiques sans erreur ; aucun HTTP 500",
         "Aucune exception ; comportement idempotent"],
    ],
    col_widths=[1.5, 3.5, 3, 3.5, 4, 3]
)

add_heading(doc, "Script de vérification contenu PDF (Python)", 3, "2E74B5")
add_code_block(doc,
"""import pdfplumber
with pdfplumber.open("recu_candidature.pdf") as pdf:
    text = pdf.pages[0].extract_text()
    assert "Ahmed Ben Ali" in text, "Nom absent du PDF !"
    assert "ISIMM" in text, "Logo/mention ISIMM absente !"
    assert len(pdf.pages[0].images) > 0, "QR Code absent !"
    print("✅ Contenu PDF validé")""")

doc.add_page_break()

# ── SECTION 1.3 — QR CODE ─────────────────────────────────────────────────────

add_heading(doc, "1.3 — Sécurité QR Code (SHA-256)", 2)
add_paragraph(doc,
    "Flux attendu : PDF contient un QR Code → Le QR Code encode une URL ou un hash SHA-256 "
    "→ L'URL de vérification retourne les données candidature si le hash est valide.", italic=True)
add_paragraph(doc, "Endpoint cible : GET /api/candidatures/verifier-qr/?token={sha256_hash}", bold=True)
doc.add_paragraph()

make_table(doc,
    ["ID", "Scénario", "Inputs", "Actions", "Résultat attendu", "Vérification DB"],
    [
        ["QR-01", "Scan QR valide",
         "PDF généré en PDF-01",
         "1. Scanner le QR Code avec smartphone (Google Lens)\n2. Ouvrir l'URL scannée",
         "Page web : N° candidature, Nom, Parcours, Date, mention ✅ Candidature authentifiée",
         "CandidatureToken.objects.get(token=hash) — is_valid=True"],
        ["QR-02", "Hash SHA-256 — intégrité",
         "Hash extrait du QR Code",
         "1. Recalculer le hash SHA-256 côté Django\n2. Comparer avec le hash du QR",
         "Hash identique (64 caractères hexadécimaux)",
         "hashlib.sha256(...).hexdigest() == hash_qr"],
        ["QR-03", "Token falsifié",
         "URL du QR avec 1 caractère modifié manuellement",
         "Ouvrir l'URL falsifiée dans le navigateur",
         "Page : ❌ QR Code invalide ou expiré — Accès refusé ; HTTP 404 ou 400",
         "Aucune entrée trouvée ; tentative loguée"],
        ["QR-04", "Token expiré",
         "Token créé il y a 30 jours (si expiration configurée)",
         "Appeler l'URL de vérification avec le vieux token",
         "❌ Ce QR Code a expiré ; HTTP 410 Gone",
         "is_expired=True sur l'objet token"],
        ["QR-05", "QR lisible à l'impression",
         "PDF imprimé N&B, 600 DPI",
         "Scanner le QR imprimé avec 3 smartphones différents",
         "QR scanné avec succès sur les 3 appareils en < 3 secondes",
         "—"],
        ["QR-06", "Unicité des tokens",
         "2 candidatures différentes",
         "Générer les PDFs des 2 candidatures",
         "Les 2 QR Codes encodent 2 hashs distincts (aucune collision)",
         "Token A != Token B en DB"],
    ],
    col_widths=[1.5, 3, 3, 3.5, 4, 3.5]
)

add_heading(doc, "Script de vérification QR Code (Python)", 3, "2E74B5")
add_code_block(doc,
"""import hashlib, requests

token_from_qr = "abc123..."  # valeur scannée depuis le PDF

# Recalculer depuis la DB
from candidature_app.models import Candidature
c = Candidature.objects.get(numero="2603-00001-GL")
expected = hashlib.sha256(
    f"{c.id}{c.numero}{c.date_soumission}".encode()
).hexdigest()

assert token_from_qr == expected, "❌ Hash invalide !"
print("✅ QR Code authentique")

# Vérifier l'endpoint
r = requests.get(f"http://localhost:8000/api/candidatures/verifier-qr/?token={token_from_qr}")
assert r.status_code == 200, f"HTTP {r.status_code}"
print("✅ Endpoint QR valide")""")

doc.add_page_break()

# ── SECTION 2 — US ────────────────────────────────────────────────────────────

add_heading(doc, "SECTION 2 — Parcours End-to-End par User Story", 1)

add_heading(doc, "Données de test globales", 2)
make_table(doc,
    ["Rôle", "Identifiants", "Données profil"],
    [
        ["Admin", "admin@isimm.tn / Admin@1234", "Superuser Django — accès total"],
        ["Responsable", "responsable@isimm.tn / Resp@1234", "Groupe commission_responsable"],
        ["Candidat 1", "ahmed.benali@test.tn / Cand@1234", "Ahmed Ben Ali, CIN 12345678, Licence Info ISIMM"],
        ["Candidat 2", "sonia.trabelsi@test.tn / Cand@5678", "Sonia Trabelsi, CIN 87654321, Licence EEA Externe"],
    ],
    col_widths=[4, 6, 8]
)

# US-01
add_heading(doc, "US-01 — Admin configure les parcours et les rôles", 2)
add_paragraph(doc, "URL : http://localhost:4200/admin/dashboard", bold=True)
doc.add_paragraph()
make_table(doc,
    ["Étape", "Action Angular", "Données saisies", "Résultat attendu Angular", "Résultat attendu DB"],
    [
        ["1", "Se connecter en Admin", "admin@isimm.tn / Admin@1234",
         "Redirection /admin/dashboard ; sidebar Admin visible",
         "request.user.is_superuser = True"],
        ["2", "Naviguer ?view=parcours-master", "—",
         "Tableau : 5 parcours Master (MPGL, MPDS, MP3I, MRGL, MRMI) — statut OUVERT",
         "SpecialiteParcoursMapping.objects.filter(type_formation='master').count() == 5"],
        ["3", "Naviguer ?view=parcours-ingenieur", "—",
         "Tableau : parcours ING_GL (Cycle Ingénieur)",
         "SpecialiteParcoursMapping.objects.filter(type_formation='cycle_ingenieur').count() == 1"],
        ["4", "Naviguer ?view=listes-selection", "—",
         "Section \"Listes de sélection\" visible",
         "—"],
        ["5", "Attribuer le rôle Responsable", "Sélectionner \"Responsable Commission\"",
         "Toast ✅ Rôle attribué",
         "User.groups contient commission_responsable pour responsable@isimm.tn"],
        ["6", "Naviguer ?view=users", "—",
         "Liste utilisateurs ; Ahmed Ben Ali et Sonia Trabelsi avec rôle Candidat",
         "User.objects.filter(groups__name='candidat').count() >= 2"],
        ["7", "Vérifier absence bouton Ajouter parcours", "—",
         "Aucun bouton + Ajouter dans la vue parcours-master",
         "—"],
    ],
    col_widths=[1.2, 4, 3, 4.5, 5.5]
)

doc.add_page_break()

# US-02
add_heading(doc, "US-02 — Responsable synchronise et ouvre les offres", 2)
add_paragraph(doc, "URL : http://localhost:4200/commission/dashboard", bold=True)
doc.add_paragraph()
make_table(doc,
    ["Étape", "Action Angular", "Données saisies", "Résultat attendu Angular", "Résultat attendu DB"],
    [
        ["1", "Connexion Responsable", "responsable@isimm.tn / Resp@1234",
         "Dashboard Commission ; menu Offre De Préinscription visible",
         "request.user dans groupe commission_responsable"],
        ["2", "Cliquer Offre De Préinscription", "—",
         "Tableau avec 6 parcours : MPGL, MPDS, MP3I, MRGL, MRMI, ING-GL",
         "offresPreinscription peuplé à 6 entrées"],
        ["3", "Vérifier données de MRGL", "—",
         "Ligne MRGL : sous-type recherche, places 111, date 22/07/2026, statut Ouvert",
         "SpecialiteParcoursMapping.objects.get(code_parcours='MRGL').actif == True"],
        ["4", "Toggle STATUT MPGL : Fermée → Ouverte", "Clic sur le toggle STATUT",
         "Toggle bleu ; badge Ouverte ; toast ✅ Statut mis à jour",
         "ConfigurationAppel mis à jour : actif=True pour MPGL"],
        ["5", "Cliquer Actualiser", "—",
         "Tableau rechargé ; MPGL affiche Ouverte",
         "Idem état DB"],
        ["6", "Cliquer Éditer sur MP3I", "—",
         "Formulaire d'édition ouvert avec données MP3I pré-remplies",
         "—"],
        ["7", "Modifier capacité MP3I", "capacite_total : 30",
         "Champ mis à jour",
         "ConfigurationAppel.capacite_accueil == 30 pour MP3I"],
        ["8", "Cliquer Enregistrer", "—",
         "Toast ✅ Offre mise à jour ; retour au tableau",
         "date_mise_a_jour actualisée"],
        ["9", "Naviguer Candidatures Master", "—",
         "Tableau de candidatures vide (aucune candidature encore)",
         "Candidature.objects.count() == 0"],
        ["10", "Vérifier absence du bloc Download", "—",
         "Aucun sélecteur de format d'export dans les vues parcours",
         "—"],
    ],
    col_widths=[1.2, 4, 3, 4.5, 5.5]
)

doc.add_page_break()

# US-03
add_heading(doc, "US-03 — Candidat 1 postule à MRGL (Master Recherche Génie Logiciel)", 2)
add_paragraph(doc, "URL : http://localhost:4200/candidat/dashboard", bold=True)
doc.add_paragraph()
make_table(doc,
    ["Étape", "Action Angular", "Données saisies", "Résultat attendu Angular", "Résultat attendu DB"],
    [
        ["1", "Connexion Candidat 1", "ahmed.benali@test.tn / Cand@1234",
         "Dashboard Candidat ; header 🎓 Espace Candidat — Ahmed Ben Ali",
         "Session valide"],
        ["2", "Vérifier le tableau de bord", "—",
         "Compteurs : Candidatures=0, En attente=0, Dossiers=0, Notifications=0",
         "Candidature.objects.filter(candidat=user).count() == 0"],
        ["3", "Cliquer Préinscription (sidebar)", "—",
         "Page : 6 cartes (2 Masters Recherche + 3 Masters Pro + 1 Cycle Ingénieur) ; badge OUVERT visible",
         "offresInscription.length == 6"],
        ["4", "Vérifier la carte MRGL", "—",
         "Titre: Mastere Recherche en Genie logiciel (MRGL) ; Date: 22/07/2026 ; Places: 111 ; Badge OUVERT",
         "—"],
        ["5", "Cliquer Détail sur MRGL", "—",
         "Modal : badge MR-GL ; tableau capacités ; formule de score ; liste dossier requis",
         "—"],
        ["6", "Fermer le modal", "—",
         "Modal se ferme proprement ; aucun blocage",
         "—"],
        ["7", "Cliquer Postuler sur MRGL", "—",
         "Wizard de soumission s'ouvre — Étape 1 : Informations Personnelles",
         "—"],
        ["8", "Remplir Étape 1",
         "Nom: Ben Ali / Prénom: Ahmed / CIN: 12345678 / Naissance: 15/06/2000 / Email: ahmed.benali@test.tn / Tél: 21698123456",
         "Champs validés (bordures vertes) ; bouton Suivant activé",
         "—"],
        ["9", "Cliquer Suivant — Étape 2", "—",
         "Étape 2 : Diplôme & Formation affichée",
         "—"],
        ["10", "Remplir Étape 2",
         "Nature: Licence / Spécialité: Informatique / Année: 2025 / M1: 13.5 / Sess1: Principale / M2: 14.2 / M3: 15.0 / Redoublement: 0",
         "Score calculé affiché en temps réel (ex: 18.7)",
         "Appel API /api/candidatures/calculer-score/ — score retourné"],
        ["11", "Cliquer Suivant — Étape 3", "—",
         "Étape 3 : Validation & Synthèse — récapitulatif affiché",
         "—"],
        ["12", "Cocher la déclaration", "✓ Je certifie l'exactitude des informations",
         "Bouton Soumettre ma candidature activé",
         "—"],
        ["13", "Cliquer Soumettre", "—",
         "Toast ✅ Candidature soumise ; N° affiché ex: 2603-00001-GL ; modal se ferme",
         "Candidature créée : statut='soumis' ; specialite='MRGL'"],
        ["14", "Vérifier le tableau de bord", "—",
         "Compteur Candidatures passe à 1 ; carte candidature visible dans Mes Candidatures",
         "Candidature.objects.filter(candidat=user).count() == 1"],
        ["15", "Cliquer Postuler MRGL une 2ème fois", "—",
         "Bouton devient Déjà candidaté (grisé, désactivé)",
         "dejaCandidature() == True"],
    ],
    col_widths=[1.2, 3.5, 4, 4.5, 5]
)

doc.add_page_break()

# US-04
add_heading(doc, "US-04 — Candidat 1 dépose son dossier", 2)
make_table(doc,
    ["Étape", "Action Angular", "Données saisies", "Résultat attendu Angular", "Résultat attendu DB"],
    [
        ["1", "Aller dans Dossiers de candidature", "—",
         "Tableau : 1 dossier lié à MRGL ; statut Non déposé",
         "DossierCandidature.objects.filter(candidat=user).count() == 1"],
        ["2", "Cliquer Ouvrir sur le dossier", "—",
         "Section expanded ; barre complétion 0% ; 6 pièces listées",
         "—"],
        ["3", "Uploader le formulaire de candidature", "DOC-01 (PDF propre)",
         "Progression 0→100% ; Pièce 1 ✅ Déposé ; complétion 16%",
         "Document créé ; type='formulaire'"],
        ["4", "Uploader la fiche imprimée", "DOC-01",
         "Pièce 2 ✅ ; complétion 33%",
         "—"],
        ["5", "Uploader le CV", "DOC-02",
         "Pièce 3 ✅ ; complétion 50%",
         "—"],
        ["6", "Uploader les diplômes", "DOC-01",
         "Pièce 4 ✅ ; complétion 67%",
         "—"],
        ["7", "Uploader les relevés de notes", "DOC-01",
         "Pièce 5 ✅ ; complétion 83%",
         "—"],
        ["8", "Laisser pièce 6 (facultative) non déposée", "—",
         "Complétion reste à 83% — pièce facultative exclue du total",
         "documentsTotaux==5 ; documentsValides==5"],
        ["9", "Cliquer Finaliser le dossier", "—",
         "Toast ✅ Dossier finalisé ; statut passe à Déposé",
         "DossierCandidature.dossier_depose == True"],
    ],
    col_widths=[1.2, 4, 3, 5, 5]
)

doc.add_page_break()

# US-05
add_heading(doc, "US-05 — Candidat 2 postule à MP3I (Master Professionnel 3I)", 2)
make_table(doc,
    ["Étape", "Action Angular", "Données saisies", "Résultat attendu Angular", "Résultat attendu DB"],
    [
        ["1", "Connexion Candidat 2", "sonia.trabelsi@test.tn / Cand@5678",
         "Dashboard Candidat 2 — Sonia Trabelsi",
         "Session valide"],
        ["2", "Postuler MP3I — même flux US-03",
         "Étape 2 : M1:12.0 / M2:13.0 / M3:11.5 / Redoubl:1",
         "Candidature soumise ; N°: 2603-00002-3I",
         "Candidature statut='soumis' pour Sonia ; specialite='MP3I'"],
        ["3", "Vérifier isolation des données", "—",
         "Dashboard Sonia n'affiche que sa propre candidature (pas celle d'Ahmed)",
         "Aucune fuite de données inter-utilisateurs"],
    ],
    col_widths=[1.2, 4, 3, 5, 5]
)

doc.add_paragraph()

# US-06
add_heading(doc, "US-06 — Responsable traite les candidatures", 2)
add_paragraph(doc, "URL : http://localhost:4200/commission/dashboard", bold=True)
doc.add_paragraph()
make_table(doc,
    ["Étape", "Action Angular", "Données saisies", "Résultat attendu Angular", "Résultat attendu DB"],
    [
        ["1", "Naviguer Candidatures Master", "—",
         "Tableau : 2 candidatures (Ahmed MRGL + Sonia MP3I) ; scores affichés",
         "Candidature.objects.count() == 2"],
        ["2", "Filtrer par parcours MRGL", "Sélectionner MRGL dans le filtre",
         "Seule la candidature d'Ahmed visible",
         "—"],
        ["3", "Changer statut Ahmed → Présélectionné",
         "Sélectionner preselectionne dans le dropdown",
         "Toast ✅ Statut mis à jour ; ligne mise à jour en temps réel",
         "Candidature.objects.get(id=1).statut == 'preselectionne'"],
        ["4", "Vérifier notification Ahmed", "— (basculer sur compte Ahmed)",
         "Notification : Votre candidature MRGL a été présélectionnée",
         "NotificationCandidat.objects.filter(candidat=ahmed).exists()"],
        ["5", "Exporter la liste", "Cliquer Exporter CSV",
         "Fichier CSV téléchargé avec en-têtes et données",
         "—"],
        ["6", "Rejeter Sonia avec motif", "motif: Moyenne insuffisante",
         "Toast ✅ Décision enregistrée",
         "statut='rejete' ; motif_rejet='Moyenne insuffisante'"],
    ],
    col_widths=[1.2, 4, 3, 5, 5]
)

doc.add_page_break()

# US-07
add_heading(doc, "US-07 — Candidat 1 suit sa candidature et télécharge le reçu PDF", 2)
make_table(doc,
    ["Étape", "Action Angular", "Données saisies", "Résultat attendu Angular", "Résultat attendu DB"],
    [
        ["1", "Ahmed se reconnecte", "ahmed.benali@test.tn / Cand@1234",
         "Dashboard ; Compteur Présélectionnées = 1",
         "—"],
        ["2", "Aller dans Candidatures", "—",
         "Badge statut MRGL : Présélectionné (vert) ; stepper à l'étape correcte",
         "statut = 'preselectionne'"],
        ["3", "Aller dans Suivi de candidature", "—",
         "Timeline/workflow visible ; étape Présélectionné mise en évidence",
         "—"],
        ["4", "Cliquer Récapitulatif", "—",
         "Modal récapitulatif avec toutes les données de la candidature",
         "—"],
        ["5", "Cliquer Télécharger PDF", "—",
         "PDF téléchargé (>=50 Ko) ; contient Ahmed Ben Ali, MRGL, N°, date, QR Code",
         "HTTP 200 ; Content-Type: application/pdf"],
        ["6", "Vérifier le QR Code du PDF", "Scanner avec smartphone",
         "URL s'ouvre : ✅ Candidature authentifiée — Ahmed Ben Ali — MRGL",
         "Token valide en DB"],
        ["7", "Aller dans Notifications", "—",
         "1+ notifications ; Présélectionné pour MRGL",
         "notificationsNonLues >= 1"],
        ["8", "Marquer toutes comme lues", "Clic Tout marquer comme lu",
         "Compteur badge passe à 0",
         "NotificationCandidat.lue == True pour toutes"],
    ],
    col_widths=[1.2, 4, 3, 5, 5]
)

doc.add_page_break()

# ── SECTION 3 — RÉGRESSION ────────────────────────────────────────────────────

add_heading(doc, "SECTION 3 — Tests de Régression Critiques", 1)

add_heading(doc, "3.1 — Isolation & Sécurité", 2)
make_table(doc,
    ["ID", "Test", "Action", "Résultat attendu"],
    [
        ["SEC-01", "Accès non authentifié",
         "GET /api/candidatures/mes-candidatures/ sans token",
         "HTTP 401 Unauthorized"],
        ["SEC-02", "Accès croisé entre candidats",
         "Ahmed appelle GET /api/candidatures/{id_de_sonia}/",
         "HTTP 403 ou données vides"],
        ["SEC-03", "Candidat accède à un endpoint Responsable",
         "POST /api/candidatures/changer-statut/ avec token Candidat",
         "HTTP 403 Forbidden"],
        ["SEC-04", "Token JWT expiré",
         "Modifier l'expiry en dev → faire une requête API",
         "HTTP 401 ; redirection vers /login côté Angular"],
        ["SEC-05", "Injection SQL via formulaire",
         "Saisir '; DROP TABLE candidature; -- dans le champ Nom",
         "Valeur stockée littéralement ; aucune erreur DB ; pas d'exécution SQL"],
    ],
    col_widths=[1.5, 4, 5, 7]
)

add_heading(doc, "3.2 — Cohérence des données", 2)
make_table(doc,
    ["ID", "Test", "Vérification"],
    [
        ["COH-01", "N° candidature unique",
         "Candidature.objects.values('numero').annotate(c=Count('id')).filter(c__gt=1) → résultat vide"],
        ["COH-02", "Score cohérent",
         "Score affiché Angular == score retourné par GET /api/candidatures/{id}/score/"],
        ["COH-03", "Places disponibles décrémentées",
         "Après soumission candidature MRGL : vérifier places_disponibles dans DB"],
        ["COH-04", "Date limite respectée",
         "Appeler Postuler après date_limite → HTTP 400 \"Délai de candidature dépassé\""],
    ],
    col_widths=[1.5, 5, 12]
)

add_heading(doc, "3.3 — WebSocket temps réel", 2)
make_table(doc,
    ["ID", "Test", "Action", "Résultat attendu"],
    [
        ["WS-01", "Mise à jour statut en temps réel",
         "Responsable change statut Ahmed → Observer dashboard Ahmed sans refresh",
         "Badge statut mis à jour en < 5 secondes sans rechargement de page"],
        ["WS-02", "Reconnexion automatique",
         "Couper la connexion réseau 10 secondes → La rétablir",
         "Indicateur WebSocket repasse à online ; données rechargées automatiquement"],
        ["WS-03", "Indicateur de connexion",
         "Observer l'indicateur vert/orange/rouge dans la sidebar candidat",
         "Vert = connecté ; Orange = reconnexion ; Rouge = déconnecté"],
    ],
    col_widths=[1.5, 4, 5.5, 7]
)

doc.add_page_break()

# ── SECTION 4 — MATRICE ───────────────────────────────────────────────────────

add_heading(doc, "SECTION 4 — Matrice de Validation Finale", 1)
add_paragraph(doc, "Notation : ✅ OK  |  ❌ KO  |  ⚠️ Partiel  |  ⬜ À tester", bold=True, italic=True)
doc.add_paragraph()

make_table(doc,
    ["Module", "Scénario nominal", "Cas d'erreur", "Performance", "Sécurité", "Statut"],
    [
        ["Auth JWT", "Login / Logout", "Token expiré SEC-04", "< 500 ms", "SEC-01 à 05", "⬜ À tester"],
        ["Admin — Parcours", "US-01 (7 étapes)", "API vide → fallback 6 parcours", "< 1 s", "Accès restreint Admin", "⬜ À tester"],
        ["Responsable — Offres", "US-02 (10 étapes)", "API vide → canonical 6", "< 1 s", "SEC-03", "⬜ À tester"],
        ["Candidat — Préinscription", "US-03 (15 étapes)", "API < 6 → fallback canonical", "< 1 s", "SEC-02", "⬜ À tester"],
        ["Dépôt dossier", "US-04 (9 étapes)", "OCR-03 / OCR-04", "Upload < 5 s", "SEC-02", "⬜ À tester"],
        ["OCR PaddleOCR", "OCR-01 (extraction propre)", "OCR-03 / OCR-04", "< 10 s/page", "—", "⬜ À tester"],
        ["PDF ReportLab", "PDF-01 / PDF-02 / PDF-03", "PDF-04 (données vides)", "< 3 s", "—", "⬜ À tester"],
        ["QR Code SHA-256", "QR-01 / QR-02", "QR-03 / QR-04", "Scan < 3 s", "QR-03 (falsification)", "⬜ À tester"],
        ["WebSocket", "WS-01 (temps réel)", "WS-02 (reconnexion)", "Délai < 5 s", "—", "⬜ À tester"],
        ["Notifications", "US-07 (étapes 7-8)", "API down → liste vide", "< 1 s", "SEC-02", "⬜ À tester"],
        ["Multi-candidats", "US-05 (isolation)", "Accès croisé SEC-02", "< 1 s", "SEC-02", "⬜ À tester"],
        ["Traitement Responsable", "US-06 (6 étapes)", "COH-04 (date dépassée)", "< 1 s", "SEC-03", "⬜ À tester"],
    ],
    col_widths=[3.5, 3.5, 3.5, 2.5, 2.5, 2.5]
)

doc.add_page_break()

# ── SECTION 5 — COMMANDES DB ──────────────────────────────────────────────────

add_heading(doc, "SECTION 5 — Commandes de Vérification DB Rapides", 1)
add_paragraph(doc, "Lancer depuis le dossier du service concerné :", bold=True)
add_code_block(doc, "python manage.py shell")
doc.add_paragraph()

add_heading(doc, "Résumé global de la base de données", 2)
add_code_block(doc,
"""from candidature_app.models import Candidature, SpecialiteParcoursMapping
from django.contrib.auth import get_user_model
User = get_user_model()

print("=== RÉSUMÉ GLOBAL ===")
print("Candidatures total         :", Candidature.objects.count())
print("Candidatures soumises      :", Candidature.objects.filter(statut='soumis').count())
print("Présélectionnés            :", Candidature.objects.filter(statut='preselectionne').count())
print("Rejetés                    :", Candidature.objects.filter(statut='rejete').count())
print("Candidats (utilisateurs)   :", User.objects.filter(groups__name='candidat').count())

print("\\n=== PARCOURS ACTIFS ===")
for p in SpecialiteParcoursMapping.objects.filter(actif=True).order_by('ordre'):
    print(f"  {p.code_parcours:8s} | {p.type_formation:18s} | {p.nom_parcours}")""")

add_heading(doc, "Vérification par candidat", 2)
add_code_block(doc,
"""from candidature_app.models import Candidature
from django.contrib.auth import get_user_model
User = get_user_model()

ahmed = User.objects.get(email='ahmed.benali@test.tn')
cands = Candidature.objects.filter(candidat=ahmed)
print(f"Ahmed : {cands.count()} candidature(s)")
for c in cands:
    print(f"  {c.numero} | {c.specialite} | {c.statut} | score={c.score}")""")

add_heading(doc, "Vérification isolation des données", 2)
add_code_block(doc,
"""# Vérifier qu'aucune candidature n'est visible cross-utilisateurs
from candidature_app.models import Candidature
from django.db.models import Count

# Chercher des doublons de numéro (doit être vide)
doublons = Candidature.objects.values('numero').annotate(c=Count('id')).filter(c__gt=1)
print("Doublons de N° candidature :", list(doublons))  # Attendu : []

# Tokens QR uniques
from candidature_app.models import CandidatureToken  # si le modèle existe
doublons_token = CandidatureToken.objects.values('token').annotate(c=Count('id')).filter(c__gt=1)
print("Tokens QR dupliqués        :", list(doublons_token))  # Attendu : []""")

add_heading(doc, "Vérification des 6 parcours en DB", 2)
add_code_block(doc,
"""from candidature_app.models import SpecialiteParcoursMapping

CODES_ATTENDUS = ['MPGL', 'MPDS', 'MP3I', 'MRGL', 'MRMI', 'ING_GL']
for code in CODES_ATTENDUS:
    try:
        p = SpecialiteParcoursMapping.objects.get(code_parcours=code, actif=True)
        print(f"  ✅ {code:8s} — {p.nom_parcours}")
    except SpecialiteParcoursMapping.DoesNotExist:
        print(f"  ❌ {code:8s} — ABSENT de la base !")""")

doc.add_paragraph()
add_paragraph(doc,
    "Ce document a été généré automatiquement le " +
    datetime.date.today().strftime("%d/%m/%Y") +
    " pour la validation de la Plateforme ISIMM — Gestion des Préinscriptions 2025-2026.",
    italic=True)

# ── SAVE ─────────────────────────────────────────────────────────────────────

output_path = r"C:\Users\HP\Desktop\PFE\Cahier_de_Test_ISIMM_2025_2026.docx"
doc.save(output_path)
print(f"✅ Document généré : {output_path}")
